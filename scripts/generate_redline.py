#!/usr/bin/env python3
"""
generate_redline.py — Apply Word track-changes to NDA source document in-place.

Opens the source DOCX and for each issue locates the original text in paragraphs,
then replaces it with w:del + w:ins XML so the counterparty can accept/reject each
change in Word's Review mode.

Usage:
  python3 generate_redline.py --source path/to/nda.docx --issues issues.json
  python3 generate_redline.py --source path/to/nda.docx --all-standard

Environment:
  NDA_SKILL_REVIEWER_NAME  — Author name shown in track changes (default: "Reviewer")
"""

import argparse
import json
import os
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

except ImportError:
    print("ERROR: python-docx and lxml are required. Run: pip install python-docx lxml")
    sys.exit(1)

# ── Config ─────────────────────────────────────────────────────────────────────

AUTHOR = os.environ.get("NDA_SKILL_REVIEWER_NAME", "Reviewer")
DATE_STR = f"{date.today().isoformat()}T00:00:00Z"
_id_counter = [0]

# Tags of content-bearing paragraph children (removed when rebuilding a paragraph)
_CONTENT_TAGS = None


def _nid() -> str:
    _id_counter[0] += 1
    return str(_id_counter[0])


def _content_tags() -> set:
    global _CONTENT_TAGS
    if _CONTENT_TAGS is None:
        _CONTENT_TAGS = {qn("w:r"), qn("w:ins"), qn("w:del"), qn("w:hyperlink")}
    return _CONTENT_TAGS


# ── XML helpers ────────────────────────────────────────────────────────────────

def _mk_run(text: str):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _mk_del(text: str):
    d = OxmlElement("w:del")
    d.set(qn("w:id"), _nid())
    d.set(qn("w:author"), AUTHOR)
    d.set(qn("w:date"), DATE_STR)
    r = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    r.append(dt)
    d.append(r)
    return d


def _mk_ins(text: str):
    i = OxmlElement("w:ins")
    i.set(qn("w:id"), _nid())
    i.set(qn("w:author"), AUTHOR)
    i.set(qn("w:date"), DATE_STR)
    i.append(_mk_run(text))
    return i


# ── Standard redline positions ─────────────────────────────────────────────────

STANDARD_REDLINES = [
    {
        "id": "ci-temporal",
        "clause": "CI Definition — Temporal Scope",
        "priority": "HIGH",
        "original": "whether before or after the date of this Agreement",
        "redlined": "on or after the date of this Agreement",
        "rationale": (
            "Pre-signing disclosures are typically governed by a separate NDA or implied "
            "confidentiality obligations. Including them here creates unbounded historical liability."
        ),
    },
    {
        "id": "ci-oral",
        "clause": "CI Definition — Oral Disclosures",
        "priority": "HIGH",
        "original": "",
        "redlined": (
            "For oral disclosures, Confidential Information shall include only information "
            "designated as confidential at the time of disclosure and confirmed in writing "
            "within ten (10) business days thereafter."
        ),
        "rationale": (
            "Without a written confirmation window, any oral statement could retrospectively "
            "become Confidential Information, creating practical enforcement issues."
        ),
    },
    {
        "id": "permitted-advisors",
        "clause": "Permitted Disclosees — Professional Advisors",
        "priority": "HIGH",
        "original": "directors, officers, and employees",
        "redlined": (
            "directors, officers, employees, and professional advisors "
            "(including legal counsel, accountants, and financial advisors) who are "
            "bound by professional obligations of confidentiality"
        ),
        "rationale": (
            "Professional advisors routinely need access to evaluate a transaction. "
            "Excluding them creates operational friction and is non-standard."
        ),
    },
    {
        "id": "return-destruction-election",
        "clause": "Return / Destruction — Party Election",
        "priority": "MEDIUM",
        "original": "with the Disclosing Party's written consent, will either (a) return",
        "redlined": "at the Receiving Party's election, either (a) return",
        "rationale": (
            "Requiring Disclosing Party consent for destruction removes the Receiving Party's "
            "ability to meet its own data-management obligations. Market standard is election."
        ),
    },
    {
        "id": "return-retention-exception",
        "clause": "Return / Destruction — Retention Exception",
        "priority": "MEDIUM",
        "original": "",
        "redlined": (
            "Notwithstanding the foregoing, the Receiving Party shall not be required to "
            "return or destroy copies of Confidential Information held in automated backup "
            "systems, provided such copies are not accessible in the normal course of "
            "business and are overwritten in the Receiving Party's normal backup cycle."
        ),
        "rationale": (
            "Without this carveout, complete return/destruction is technically impossible "
            "and creates residual legal exposure for the Receiving Party."
        ),
    },
    {
        "id": "term",
        "clause": "Confidentiality Tail — Term",
        "priority": "MEDIUM",
        "original": "three (3) years",
        "redlined": "two (2) years",
        "rationale": (
            "Two years is market standard for general commercial NDAs. Three years may be "
            "appropriate for highly sensitive technical information — adjust based on context."
        ),
    },
]


# ── In-place track change application ─────────────────────────────────────────

def _para_full_text(para) -> str:
    """Concatenate text from all direct runs in a paragraph."""
    return "".join(run.text for run in para.runs)


def _apply_change_to_para(para, original: str, redlined: str) -> bool:
    """
    If para contains original, rebuild paragraph XML with track changes in-place:
      pre-match text → w:del(original) → w:ins(redlined) → post-match text

    Content-bearing children (w:r, w:ins, w:del, w:hyperlink) are replaced;
    structural elements (w:pPr, w:bookmarkStart, w:bookmarkEnd) are preserved.

    Returns True if applied, False if original not found in this paragraph.
    """
    full_text = _para_full_text(para)
    if original not in full_text:
        return False

    idx = full_text.index(original)
    pre = full_text[:idx]
    post = full_text[idx + len(original):]

    p_elem = para._p
    tags = _content_tags()

    # Remove content-bearing children; keep pPr, bookmarks, sectPr, etc.
    for child in list(p_elem):
        if child.tag in tags:
            p_elem.remove(child)

    if pre:
        p_elem.append(_mk_run(pre))
    if original:
        p_elem.append(_mk_del(original))
    if redlined:
        p_elem.append(_mk_ins(redlined))
    if post:
        p_elem.append(_mk_run(post))

    return True


def _iter_paragraphs(doc) -> list:
    """Return all paragraphs from body and table cells, deduplicated.

    Merged table cells share the same underlying XML node; deduplicating by
    id(para._p) prevents the same change being applied twice.
    """
    seen: set = set()
    paras: list = []

    def _add(para):
        key = id(para._p)
        if key not in seen:
            seen.add(key)
            paras.append(para)

    for para in doc.paragraphs:
        _add(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _add(para)
    return paras


def _is_placeholder_original(original: str) -> bool:
    """True if original is empty or a descriptive placeholder (not literal NDA text)."""
    return not original or original.strip().startswith("(")


# ── Redline document generation ────────────────────────────────────────────────

def generate_redline_doc(
    source_path: Path,
    issues: list,
    output_path: Path = None,
) -> Path:
    """
    Open source DOCX and apply track changes in-place.

    For each issue:
    - Searches all paragraphs for issue['original'] (exact substring match)
    - Replaces matched text with w:del + w:ins XML
    - Issues where original is empty or a placeholder, or where the text is not
      found, are collected and appended as a [MANUAL REDLINE NEEDED] section.

    issues: list of dicts with keys: id, clause, priority, original, redlined, rationale
    """
    try:
        doc = Document(str(source_path))
    except Exception as e:
        print(f"ERROR: Could not open DOCX file {source_path}: {e}", file=sys.stderr)
        sys.exit(1)
    today = date.today().strftime("%Y%m%d")
    reviewer = AUTHOR.replace(" ", "_")

    if output_path is None:
        output_path = source_path.parent / f"NDA_Redlined_{reviewer}_{today}.docx"

    all_paras = _iter_paragraphs(doc)
    unmatched = []
    matched_count = 0

    for issue in issues:
        original = issue.get("original", "")
        redlined = issue.get("redlined", "")

        if _is_placeholder_original(original):
            # Pure insertion — no text to locate, flag for manual placement
            unmatched.append(issue)
            continue

        applied = False
        for para in all_paras:
            if _apply_change_to_para(para, original, redlined):
                applied = True
                matched_count += 1
                break  # apply to first matching paragraph only

        if not applied:
            unmatched.append(issue)

    # ── Fallback section for unmatched / pure-insertion changes ───────────────
    if unmatched:
        doc.add_paragraph("─" * 60)

        h = doc.add_paragraph()
        h.add_run("[MANUAL REDLINE NEEDED]").bold = True

        doc.add_paragraph(
            f"The following {len(unmatched)} change(s) could not be automatically "
            "located in the document text. Please apply them manually:"
        )

        for issue in unmatched:
            p = doc.add_paragraph()
            p.add_run(f"[{issue.get('priority', '?')}] {issue.get('clause', '')}: ").bold = True
            orig = issue.get("original", "")
            if orig and not _is_placeholder_original(orig):
                p._p.append(_mk_del(orig))
                p._p.append(_mk_run(" → "))
            p._p.append(_mk_ins(issue.get("redlined", "")))

            rationale = issue.get("rationale", "")
            if rationale:
                doc.add_paragraph(f"  Rationale: {rationale}")

    try:
        doc.save(str(output_path))
    except OSError as e:
        print(f"ERROR: Could not save output file {output_path}: {e}", file=sys.stderr)
        sys.exit(1)

    return output_path


# ── CLI ────────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Apply NDA track-changes to source DOCX in-place"
    )
    parser.add_argument("--source",       required=True, help="Path to source NDA file (.docx)")
    parser.add_argument("--issues",       default="",    help="Path to JSON issues file")
    parser.add_argument("--output",       default="",    help="Output DOCX path (auto-named if omitted)")
    parser.add_argument("--all-standard", action="store_true",
                        help="Apply all standard redline positions")
    args = parser.parse_args()

    source = Path(args.source).expanduser()
    if not source.exists():
        print(f"ERROR: Source file not found: {source}")
        sys.exit(1)

    # Load issues
    if args.issues:
        issues_path = Path(args.issues).expanduser()
        try:
            issues = json.loads(issues_path.read_text())
        except FileNotFoundError:
            print(f"ERROR: Issues file not found: {issues_path}", file=sys.stderr)
            sys.exit(1)
        except PermissionError:
            print(f"ERROR: Permission denied reading: {issues_path}", file=sys.stderr)
            sys.exit(1)
        except UnicodeDecodeError as e:
            print(f"ERROR: Issues file is not valid UTF-8: {e}", file=sys.stderr)
            sys.exit(1)
        except json.JSONDecodeError as e:
            print(f"ERROR: Issues file is not valid JSON: {e}", file=sys.stderr)
            sys.exit(1)
        if not isinstance(issues, list):
            print(f"ERROR: Issues file must contain a JSON array, got {type(issues).__name__}", file=sys.stderr)
            sys.exit(1)
    elif args.all_standard:
        issues = STANDARD_REDLINES
    else:
        print(f"No --issues file provided. Applying {len(STANDARD_REDLINES)} standard redline positions.")
        issues = STANDARD_REDLINES

    output = Path(args.output).expanduser() if args.output else None
    out_path = generate_redline_doc(source, issues, output)

    manual_count = sum(1 for i in issues if _is_placeholder_original(i.get("original", "")))
    eligible_count = len(issues) - manual_count

    print(f"✅ Redlined document generated: {out_path}")
    print(f"   In-place eligible: {eligible_count} | Manual section: {manual_count}")
    for issue in issues:
        tag = "(manual)" if _is_placeholder_original(issue.get("original", "")) else ""
        print(f"   [{issue['priority']}] {issue['clause']} {tag}".rstrip())


if __name__ == "__main__":
    main()
